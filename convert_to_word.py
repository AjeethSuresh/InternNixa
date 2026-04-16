import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_word_report(md_path, doc_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Define styles or just set manually
    for line in lines:
        line = line.strip()
        
        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        
        # Horizontal Rule
        elif line == '---':
            p = doc.add_paragraph()
            p.add_run('__________________________________________________________________').bold = True
        
        # Equations (Double $)
        elif line.startswith('$$') and line.endswith('$$'):
            eq = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(eq)
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.italic = True

        # List items
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line[line.find('.')+1:].strip(), style='List Number')

        # Empty line
        elif not line:
            continue
        
        # Paragraphs (Normal text)
        else:
            # Handle inline bold/italic/math
            # Simple replacement for **bold** and *italic* and $math$
            p = doc.add_paragraph()
            
            # This is a bit complex for a one-off script, let's keep it simple
            # but handle common cases
            processed_line = line
            # Bold
            processed_line = re.sub(r'\*\*(.*?)\*\*', r'|BOLD|\1|END|', processed_line)
            # Italic
            processed_line = re.sub(r'\*(.*?)\*', r'|ITALIC|\1|END|', processed_line)
            # Inline math
            processed_line = re.sub(r'\$(.*?)\$', r'|MATH|\1|END|', processed_line)
            
            parts = re.split(r'(\|BOLD\||\|ITALIC\||\|MATH\||\|END\|)', processed_line)
            
            current_bold = False
            current_italic = False
            current_math = False
            
            for part in parts:
                if part == "|BOLD|": current_bold = True
                elif part == "|ITALIC|": current_italic = True
                elif part == "|MATH|": current_math = True
                elif part == "|END|":
                    current_bold = False
                    current_italic = False
                    current_math = False
                else:
                    if not part: continue
                    run = p.add_run(part)
                    if current_bold: run.bold = True
                    if current_italic: run.italic = True
                    if current_math: 
                        run.font.name = 'Consolas'
                        run.italic = True

    doc.save(doc_path)
    print(f"Word document saved to: {doc_path}")

if __name__ == "__main__":
    md_file = "TECHNICAL_REPORT_EXTENDED.md"
    doc_file = "TECHNICAL_REPORT_EXTENDED.docx"
    create_word_report(md_file, doc_file)
